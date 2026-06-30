"""
촬영 스케줄 → Word(.docx) 생성. 공유받은 듀먼 촬영 스케줄 양식과 동일한 구성:
제목 → 장소별 섹션(■ 장소 N (N컷) / 복장·셋업) → 컷 표(4컷/행, 각 컷 = 태그·화면동작·나레이션).
"""
from __future__ import annotations

import io


def build_docx(schedule: dict, meta: dict | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    meta = meta or {}
    doc = Document()
    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(9)

    def shade(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_color)
        tcPr.append(sh)

    # 제목
    title = schedule.get("title") or (meta.get("title") or "촬영 스케줄")
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(title); run.bold = True; run.font.size = Pt(18)
    sub = doc.add_paragraph()
    sline = " · ".join([x for x in [meta.get("brand"), meta.get("product"), meta.get("date")] if x]) or "촬영용 — 장소별 동선 순서 배치"
    sr = sub.add_run(sline); sr.font.size = Pt(10); sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    locations = schedule.get("locations") or []
    for li, loc in enumerate(locations):
        cuts = loc.get("cuts") or []
        # 장소 헤더
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(14)
        pr = ph.add_run(f"■ {loc.get('location','장소')}  ({len(cuts)}컷)")
        pr.bold = True; pr.font.size = Pt(12); pr.font.color.rgb = RGBColor(0x0b, 0x5c, 0xd6)
        # 복장 · 셋업
        info = doc.add_paragraph()
        ir = info.add_run(f"복장: {loc.get('wardrobe','-')}  |  셋업: {loc.get('setup','-')}")
        ir.font.size = Pt(9); ir.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # 컷 표 (4컷/행)
        if cuts:
            per = 4
            for start in range(0, len(cuts), per):
                chunk = cuts[start:start + per]
                table = doc.add_table(rows=2, cols=len(chunk))
                table.style = "Table Grid"
                table.autofit = True
                for ci, cut in enumerate(chunk):
                    # 1행: 태그
                    tcell = table.cell(0, ci); shade(tcell, "EEF3FB")
                    tp = tcell.paragraphs[0]; tr = tp.add_run(f"#{start+ci+1}  [{cut.get('tag','')}]")
                    tr.bold = True; tr.font.size = Pt(8); tr.font.color.rgb = RGBColor(0x0b, 0x5c, 0xd6)
                    # 2행: 화면 동작 + 나레이션
                    bcell = table.cell(1, ci)
                    ap = bcell.paragraphs[0]
                    ar = ap.add_run(cut.get("action", "")); ar.font.size = Pt(9)
                    np = bcell.add_paragraph()
                    nlabel = np.add_run("나레이션: "); nlabel.bold = True; nlabel.font.size = Pt(8.5)
                    nr = np.add_run(cut.get("narration", "")); nr.font.size = Pt(8.5); nr.font.color.rgb = RGBColor(0xc2, 0x41, 0x0c)
                doc.add_paragraph()
        # 복장 변경 안내
        if li + 1 < len(locations):
            nxt = locations[li + 1].get("wardrobe", "")
            if nxt and nxt != loc.get("wardrobe", ""):
                ch = doc.add_paragraph(); cr = ch.add_run(f"🔄 복장 변경: {loc.get('wardrobe','')} → {nxt}")
                cr.font.size = Pt(9); cr.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
